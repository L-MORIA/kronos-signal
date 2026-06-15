"""Конвертация trader-guide.md → trader-guide.docx
Запуск: python C:\kronos-signal\scripts\md_to_docx.py
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

MD_PATH = "C:/kronos-signal/references/trader-guide.md"
DOCX_PATH = "C:/kronos-signal/references/trader-guide.docx"

def add_code_block(doc, text):
    """Код — моноширинный, серый фон (через стиль)."""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    # серый фон через shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '#1E293B')
    shd.set(qn('w:val'), 'clear')
    run._element.get_or_add_rPr().append(shd)

def md_to_docx(md_text):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    def flush_code():
        if code_buffer:
            add_code_block(doc, "\n".join(code_buffer))
            code_buffer.clear()

    def flush_table():
        if len(table_buffer) < 2:
            table_buffer.clear()
            return
        # table_buffer[0] = header, table_buffer[1] = separator, rest = rows
        header_cells = [c.strip() for c in table_buffer[0].split("|") if c.strip()]
        rows = []
        for line in table_buffer[2:]:
            if "|" in line and not line.startswith("---"):
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    rows.append(cells)
        if not rows:
            table_buffer.clear()
            return
        table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header
        for j, h in enumerate(header_cells):
            cell = table.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
        # Data
        for ri, row in enumerate(rows):
            for j, val in enumerate(row):
                if j < len(header_cells):
                    table.rows[ri + 1].cells[j].text = val
                    for p in table.rows[ri + 1].cells[j].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
        doc.add_paragraph()  # spacing after table
        table_buffer.clear()

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table
        if "|" in line and line.count("|") >= 2:
            # Check if next line is ---| pattern
            if i + 1 < len(lines) and "---" in lines[i + 1]:
                table_buffer = [line]
                # Collect all table lines
                j = i
                while j < len(lines) and "|" in lines[j]:
                    table_buffer.append(lines[j])
                    j += 1
                flush_table()
                i = j
                continue
            else:
                if not in_table:
                    doc.add_paragraph(line.strip())
                i += 1
                continue

        in_table = False

        # Separator
        if line.startswith("---") or line.startswith("==="):
            # horizontal rule — тонкая линия
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            i += 1
            continue

        # Empty line
        if not line.strip():
            flush_table()
            i += 1
            continue

        # Headers
        if line.startswith("# "):
            flush_table()
            p = doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith("## "):
            flush_table()
            p = doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("### "):
            flush_table()
            p = doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("#### "):
            flush_table()
            p = doc.add_heading(line[5:], level=4)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            flush_table()
            text = line.strip()[2:]
            # Bold pattern: **text**
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line.strip()):
            flush_table()
            text = re.sub(r"^\d+\.\s", "", line.strip())
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        flush_table()
        p = doc.add_paragraph()
        # Parse inline formatting: **bold**, `code`
        parts = re.split(r"(\*\*.*?\*\*|`.*?`)", line.strip())
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            else:
                run = p.add_run(part)
        i += 1

    doc.save(DOCX_PATH)
    return DOCX_PATH

# ── Main ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()
    out = md_to_docx(md)
    size = os.path.getsize(out)
    print(f"Конвертация завершена:")
    print(f"  Вход:  {MD_PATH}")
    print(f"  Выход: {out}")
    print(f"  Размер: {size:,} байт")
